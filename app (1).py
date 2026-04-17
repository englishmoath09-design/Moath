import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from PIL import Image
import copy

# ─── Page Config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="تقرير الفرز من المصدر",
    page_icon="📋",
    layout="centered",
)

# ─── CSS: RTL + Styling ────────────────────────────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Cairo', sans-serif;
        direction: rtl;
    }

    .main-title {
        text-align: center;
        font-size: 2rem;
        font-weight: 700;
        color: #1a5c2a;
        padding: 1rem 0 0.5rem;
    }

    .section-title {
        font-size: 1.1rem;
        font-weight: 700;
        color: #2e7d32;
        border-right: 4px solid #2e7d32;
        padding-right: 10px;
        margin: 1.5rem 0 0.8rem;
    }

    .point-card {
        background: #f0f7f1;
        border: 1px solid #a5d6a7;
        border-radius: 10px;
        padding: 12px 16px;
        margin-bottom: 10px;
        direction: rtl;
    }

    .point-card-header {
        font-weight: 700;
        color: #1b5e20;
        margin-bottom: 4px;
    }

    .stButton > button {
        direction: rtl;
        font-family: 'Cairo', sans-serif;
        font-weight: 600;
    }

    .add-btn > button {
        background-color: #2e7d32 !important;
        color: white !important;
        border-radius: 8px;
    }

    .export-btn > button {
        background-color: #1565c0 !important;
        color: white !important;
        border-radius: 8px;
        width: 100%;
    }

    .clear-btn > button {
        background-color: #c62828 !important;
        color: white !important;
        border-radius: 8px;
    }

    /* Fix selectbox/input RTL */
    .stSelectbox > div, .stTextInput > div, .stTextArea > div {
        direction: rtl;
        text-align: right;
    }

    /* Divider */
    hr { border-color: #c8e6c9; }
</style>
""", unsafe_allow_html=True)

# ─── Session State Init ────────────────────────────────────────────────────────
if "points" not in st.session_state:
    st.session_state.points = []   # list of {num, note, image_bytes, image_name}
if "point_counter" not in st.session_state:
    st.session_state.point_counter = 1

# ─── Header ───────────────────────────────────────────────────────────────────
st.markdown('<div class="main-title">📋 تقرير مشروع الفرز من المصدر</div>', unsafe_allow_html=True)
st.markdown("---")

# ─── Section 1: Basic Info ─────────────────────────────────────────────────────
st.markdown('<div class="section-title">📅 البيانات الأساسية</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    days_ar = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]
    selected_day = st.selectbox("اليوم", days_ar, key="day_select")
with col2:
    selected_date = st.date_input("التاريخ", key="date_input", format="DD/MM/YYYY")

# ─── Section 2: Add Point ──────────────────────────────────────────────────────
st.markdown("---")
st.markdown('<div class="section-title">➕ إضافة نقطة ميدانية</div>', unsafe_allow_html=True)

with st.container():
    note_input = st.text_area(
        f"الملاحظات للنقطة رقم ({st.session_state.point_counter})",
        placeholder="اكتب الملاحظات الميدانية هنا...",
        height=100,
        key="note_input"
    )
    image_upload = st.file_uploader(
        "رفع صورة (اختياري)",
        type=["jpg", "jpeg", "png", "webp"],
        key="image_upload"
    )

    if image_upload:
        st.image(image_upload, caption="معاينة الصورة", width=250)

add_col, _ = st.columns([1, 3])
with add_col:
    st.markdown('<div class="add-btn">', unsafe_allow_html=True)
    add_clicked = st.button("➕ إضافة النقطة", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

if add_clicked:
    img_bytes = None
    img_name = None
    if image_upload:
        img_bytes = image_upload.read()
        img_name = image_upload.name

    st.session_state.points.append({
        "num": st.session_state.point_counter,
        "note": note_input.strip() if note_input else "—",
        "image_bytes": img_bytes,
        "image_name": img_name,
    })
    st.session_state.point_counter += 1
    st.success(f"✅ تمت إضافة النقطة بنجاح! (إجمالي النقاط: {len(st.session_state.points)})")
    st.rerun()

# ─── Section 3: Points Review ──────────────────────────────────────────────────
if st.session_state.points:
    st.markdown("---")
    st.markdown(f'<div class="section-title">📝 النقاط المدخلة ({len(st.session_state.points)} نقطة)</div>', unsafe_allow_html=True)

    for pt in st.session_state.points:
        with st.container():
            st.markdown(f"""
            <div class="point-card">
                <div class="point-card-header">📍 النقطة رقم {pt['num']}</div>
                <div>💬 <b>الملاحظات:</b> {pt['note']}</div>
                <div>🖼️ <b>الصورة:</b> {'مرفقة ✓' if pt['image_bytes'] else 'لا توجد صورة'}</div>
            </div>
            """, unsafe_allow_html=True)
            if pt["image_bytes"]:
                st.image(pt["image_bytes"], width=180)

    clear_col, _ = st.columns([1, 3])
    with clear_col:
        st.markdown('<div class="clear-btn">', unsafe_allow_html=True)
        if st.button("🗑️ حذف جميع النقاط", use_container_width=True):
            st.session_state.points = []
            st.session_state.point_counter = 1
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

# ─── Section 4: Export ─────────────────────────────────────────────────────────
st.markdown("---")
st.markdown('<div class="section-title">📄 تصدير التقرير</div>', unsafe_allow_html=True)

if not st.session_state.points:
    st.info("📌 أضف نقاطاً ميدانية أولاً لتتمكن من تصدير التقرير.")
else:
    st.markdown(f"**جاهز للتصدير:** {len(st.session_state.points)} نقطة — {selected_day} {selected_date.strftime('%d/%m/%Y')}")

    def set_cell_rtl(cell):
        """Force RTL on a table cell's paragraphs."""
        for para in cell.paragraphs:
            pPr = para._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.insert(0, bidi)
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'right')
            pPr.append(jc)

    def set_rtl_paragraph(para):
        pPr = para._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)

    def build_docx(points, day, date_str):
        doc = Document()

        # ── Page margins ──
        section = doc.sections[0]
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

        # ── Document-level RTL ──
        settings = doc.settings.element
        doc_defaults = settings.find(qn('w:docDefaults'))
        if doc_defaults is None:
            doc_defaults = OxmlElement('w:docDefaults')
            settings.insert(0, doc_defaults)
        rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = OxmlElement('w:rPrDefault')
            doc_defaults.append(rPrDefault)
        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDefault.append(rPr)
        rtl_elem = OxmlElement('w:rtl')
        rPr.append(rtl_elem)

        # ── Title ──
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("تقرير مشروع الفرز من المصدر")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2a)
        title_run.font.name = "Arial"

        doc.add_paragraph()  # spacer

        # ── Date / Day info ──
        info_para = doc.add_paragraph()
        set_rtl_paragraph(info_para)
        info_run = info_para.add_run(f"اليوم: {day}          التاريخ: {date_str}")
        info_run.bold = True
        info_run.font.size = Pt(12)
        info_run.font.name = "Arial"

        doc.add_paragraph()  # spacer

        # ── Table ──
        # Columns: م | الصورة | الملاحظات  (RTL: right-to-left column order)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        # Column widths (total usable width ~6.5 inches = 6240 DXA after 1" margins each side)
        col_widths = [Inches(0.5), Inches(2.3), Inches(3.7)]

        # Header row
        hdr_cells = table.rows[0].cells
        headers = ["م", "الصورة", "الملاحظات"]

        # Apply RTL column order visually (Arabic docs: right=م, middle=صورة, left=ملاحظات)
        for i, (cell, hdr_text) in enumerate(zip(hdr_cells, headers)):
            cell.width = col_widths[i]
            cell.text = hdr_text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Arial"
            # Header shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '2e7d32')
            tcPr.append(shd)
            # White text
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows
        for pt in points:
            row_cells = table.add_row().cells

            # Col 0: row number (م)
            row_cells[0].width = col_widths[0]
            num_para = row_cells[0].paragraphs[0]
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num_run = num_para.add_run(str(pt["num"]))
            num_run.font.size = Pt(11)
            num_run.font.name = "Arial"

            # Col 1: image
            row_cells[1].width = col_widths[1]
            img_para = row_cells[1].paragraphs[0]
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if pt["image_bytes"]:
                try:
                    img_io = io.BytesIO(pt["image_bytes"])
                    img_pil = Image.open(img_io)
                    # Convert RGBA / palette to RGB for JPEG compat
                    if img_pil.mode in ("RGBA", "P", "LA"):
                        img_pil = img_pil.convert("RGB")
                    # Resize to fit cell (max 2 inches wide, maintain aspect)
                    max_w_px = 192  # 2 inches @ 96dpi
                    w, h = img_pil.size
                    if w > max_w_px:
                        ratio = max_w_px / w
                        img_pil = img_pil.resize((max_w_px, int(h * ratio)), Image.LANCZOS)
                    buf = io.BytesIO()
                    img_pil.save(buf, format="PNG")
                    buf.seek(0)
                    img_run = img_para.add_run()
                    img_run.add_picture(buf, width=Inches(2.0))
                except Exception:
                    img_para.add_run("خطأ في الصورة")
            else:
                img_para.add_run("—")

            # Col 2: notes
            row_cells[2].width = col_widths[2]
            note_para = row_cells[2].paragraphs[0]
            set_cell_rtl(row_cells[2])
            note_run = note_para.add_run(pt["note"])
            note_run.font.size = Pt(11)
            note_run.font.name = "Arial"

        # Vertical alignment all cells
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ── Save to buffer ──
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    st.markdown('<div class="export-btn">', unsafe_allow_html=True)
    if st.button("📄 تجهيز وتحميل التقرير", use_container_width=True):
        date_str = selected_date.strftime("%d/%m/%Y")
        docx_buf = build_docx(st.session_state.points, selected_day, date_str)
        file_name = f"تقرير_الفرز_{selected_date.strftime('%Y%m%d')}.docx"
        st.download_button(
            label="⬇️ تحميل ملف Word",
            data=docx_buf,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.success("✅ تم تجهيز التقرير! اضغط زر التحميل أعلاه.")
    st.markdown('</div>', unsafe_allow_html=True)

# ─── Footer ───────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown('<p style="text-align:center; color:#888; font-size:0.85rem;">نظام إدارة التقارير الميدانية — مشروع الفرز من المصدر</p>', unsafe_allow_html=True)
